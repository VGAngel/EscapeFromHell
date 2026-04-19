#!/usr/bin/env python3
"""Convert GDD.md to GDD.docx with basic formatting."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def md_to_docx(md_path: str, docx_path: str):
    doc = Document()

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

        # H2
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)

        # H3
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)

        # H4
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)

        # Table
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            # skip separator row (---|---)
            rows = [r for r in table_lines if not re.match(r'^\|[-| :]+\|$', r)]
            if rows:
                cols = [c.strip() for c in rows[0].strip('|').split('|')]
                t = doc.add_table(rows=len(rows), cols=len(cols))
                t.style = 'Table Grid'
                for ri, row_line in enumerate(rows):
                    cells = [c.strip() for c in row_line.strip('|').split('|')]
                    for ci, cell_text in enumerate(cells):
                        if ci < len(t.rows[ri].cells):
                            t.rows[ri].cells[ci].text = cell_text
                            if ri == 0:
                                run = t.rows[ri].cells[ci].paragraphs[0].runs
                                if run:
                                    run[0].bold = True
                                set_cell_bg(t.rows[ri].cells[ci], 'D9E2F3')
            continue

        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.3)

        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)

        # Bullet
        elif line.startswith('- '):
            text = line[2:]
            # strip inline code and bold for simplicity
            text = re.sub(r'`([^`]+)`', r'\1', text)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')

        # Empty line
        elif line.strip() == '':
            doc.add_paragraph('')

        # Normal paragraph
        else:
            text = re.sub(r'`([^`]+)`', r'\1', line)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            if text.strip():
                doc.add_paragraph(text)

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    md_to_docx(os.path.join(base, 'GDD.md'), os.path.join(base, 'GDD.docx'))
